import streamlit as st
from PIL import Image
from datetime import datetime
import tempfile
import torch
import time
import fitz
import docx
import pdfplumber
import pytesseract
from docx import Document
from fpdf import FPDF
from dotenv import load_dotenv
import anthropic
import os

# --- Load environment variables ---
load_dotenv()

# --- Setup and Config ---
st.set_page_config(page_title="VIDA Multi-Model Chatbot", layout="wide")

# --- Sidebar Image ---
st.logo("https://logodix.com/logo/951536.png",
    size = "large",
    icon_image = "https://logodix.com/logo/951515.png"
)

st.title("VIDA Chatbot Tasks Prototype")
st.header("Multi-LLM Chatbot Interface")

device = "cuda" if torch.cuda.is_available() else "cpu"
ANTHROPIC_API_KEY = st.secrets.get("ANTHROPIC_API_KEY")

# --- Sidebar Model Selector ---
model_choice = st.sidebar.selectbox(
    "Choose a Model",
    ["SmolVLM2 - Video & Image Analysis", "Document Extractor", "Claude"]
)

# --- Save Functions ---
def save_as_docx(text, filename="generated_document.docx"):
    doc = Document()
    doc.add_paragraph(text)
    doc.save(filename)
    return filename

def save_as_pdf(text, filename="generated_document.pdf"):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    for line in text.split("\n"):
        pdf.cell(200, 10, txt=line, ln=True)
    pdf.output(filename)
    return filename

# --- Claude Setup ---
client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

def call_claude(prompt, model="claude-3-haiku-20240307"):
    response = client.messages.create(
        model=model,
        max_tokens=1024,
        system="You are a helpful assistant.",
        messages=[{"role": "user", "content": prompt}]
    )
    return response.content[0].text

# --- Text Extraction Functions ---
def extract_text_from_pdf(file):
    with pdfplumber.open(file) as pdf:
        return "\n".join(page.extract_text() or "" for page in pdf.pages)

def extract_text_from_docx(file):
    doc = docx.Document(file)
    return "\n".join(paragraph.text for paragraph in doc.paragraphs)

# --- SmolVLM2 Loader and Inference ---
@st.cache_resource
def load_smolvlm2():
    from transformers import AutoProcessor, AutoModelForImageTextToText
    processor = AutoProcessor.from_pretrained("HuggingFaceTB/SmolVLM2-256M-Video-Instruct")
    model = AutoModelForImageTextToText.from_pretrained(
        "HuggingFaceTB/SmolVLM2-256M-Video-Instruct",
        torch_dtype=torch.bfloat16 if device == "cuda" else torch.float32,
    ).to(device)
    return model, processor

def generate_response_smolvlm2(text_input=None, image_input=None, video_input=None, model=None, processor=None):
    content = []
    if video_input:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".mp4") as tmp:
            tmp.write(video_input.read())
            content.append({"type": "video", "path": tmp.name})
    if image_input:
        image = Image.open(image_input)
        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
            image.save(tmp, format="PNG")
            content.append({"type": "image", "url": tmp.name})
    if text_input:
        content.append({"type": "text", "text": text_input})
    messages = [{"role": "user", "content": content}]
    inputs = processor.apply_chat_template(
        messages,
        add_generation_prompt=True,
        tokenize=True,
        return_dict=True,
        return_tensors="pt"
    ).to(device, dtype=torch.bfloat16 if device == "cuda" else torch.float32)
    output_ids = model.generate(**inputs, do_sample=False, max_new_tokens=1024)
    return processor.batch_decode(output_ids, skip_special_tokens=True)[0]

# --- Interface and Chat ---
if "multimodal_messages" not in st.session_state:
    st.session_state["multimodal_messages"] = [{
        "role": "assistant",
        "content": "Hello! I'm here to help. Select a model and send your instructions.",
        "timestamp": datetime.now().strftime("%I:%M %p"),
    }]
if "uploader_key" not in st.session_state:
    st.session_state["uploader_key"] = 0

st.markdown("<div id='chat_start'></div>", unsafe_allow_html=True)
for msg in st.session_state["multimodal_messages"]:
    with st.chat_message(msg["role"]):
        st.markdown(msg["content"])
        st.caption(f"Sent at {msg['timestamp']}")

user_text = None if model_choice == "Document Extractor" else st.chat_input("Type your instructions here:")

# --- Model-specific inputs ---
image_file = None
video_file = None
uploaded_doc = None

if model_choice == "Claude":
    st.info("Model Selected: Claude", icon="🤖")
    st.info("⚠️ Claude currently only supports text input. Image and video uploads are hidden.")

elif model_choice == "Document Extractor":
    st.info("Model Selected: Document Extractor", icon="📄")
    st.info("Upload a DOCX or PDF file to extract text from.")
    uploaded_doc = st.file_uploader("Upload DOCX or PDF", type=["docx", "pdf"], key=f"doc_{st.session_state['uploader_key']}")
    if uploaded_doc:
        if st.button("📄 Extract Text"):
            with st.spinner("Extracting text from document..."):
                if uploaded_doc.name.endswith(".pdf"):
                    response = extract_text_from_pdf(uploaded_doc)
                elif uploaded_doc.name.endswith(".docx"):
                    response = extract_text_from_docx(uploaded_doc)
                else:
                    response = "❌ Unsupported file type. Please upload a DOCX or PDF."
            st.session_state["multimodal_messages"].append({
                "role": "assistant",
                "content": response,
                "timestamp": datetime.now().strftime("%I:%M %p")
            })
            st.session_state["uploader_key"] += 1
            st.markdown("<script>location.href='#chat_start'</script>", unsafe_allow_html=True)
            st.rerun()

elif model_choice == "SmolVLM2 - Video & Image Analysis":
    st.info("Model Selected: SmolVLM2", icon="🎬")
    image_file = st.file_uploader("Upload an image (optional)", type=["png", "jpg", "jpeg"], key=f"img_{st.session_state['uploader_key']}")
    video_file = st.file_uploader("Upload a video (optional)", type=["mp4", "avi", "mov"], key=f"vid_{st.session_state['uploader_key']}" )

if user_text and user_text.strip():
    ts = datetime.now().strftime("%I:%M %p")
    st.session_state["multimodal_messages"].append({"role": "user", "content": user_text, "timestamp": ts})

    with st.spinner("Generating response..."):
        if model_choice == "SmolVLM2 - Video & Image Analysis":
            model, processor = load_smolvlm2()
            response = generate_response_smolvlm2(user_text, image_file, video_file, model, processor)

        elif model_choice == "Claude":
            try:
                response = call_claude(user_text)
            except Exception as e:
                response = f"❌ Claude API Error: {e}"

    st.session_state["multimodal_messages"].append({
        "role": "assistant",
        "content": response,
        "timestamp": datetime.now().strftime("%I:%M %p")
    })

    st.session_state["uploader_key"] += 1
    st.rerun()

if st.button("🗑️ Clear Chat"):
    st.session_state["multimodal_messages"] = []
    st.rerun()
